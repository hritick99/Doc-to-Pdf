"""
Integration tests for the DOCX to PDF conversion service
"""
import requests
import time
import os
from io import BytesIO
import zipfile
from docx import Document

BASE_URL = "http://localhost:8000/api/v1"

def create_sample_docx(filename: str, content: str) -> bytes:
    """Create a simple DOCX file in memory"""
    doc = Document()
    doc.add_heading(filename, 0)
    doc.add_paragraph(content)
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

def create_test_zip() -> BytesIO:
    """Create a zip file with sample DOCX files"""
    zip_buffer = BytesIO()
    
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        # Create 3 sample documents
        for i in range(1, 4):
            filename = f"document_{i}.docx"
            content = f"This is test document number {i}. It contains sample content for testing."
            docx_data = create_sample_docx(filename, content)
            zip_file.writestr(filename, docx_data)
    
    zip_buffer.seek(0)
    return zip_buffer

def test_full_workflow():
    """Test the complete conversion workflow"""
    print("=" * 60)
    print("Testing DOCX to PDF Conversion Service")
    print("=" * 60)
    
    # Step 1: Create and upload job
    print("\n1. Creating test documents and uploading...")
    zip_data = create_test_zip()
    
    response = requests.post(
        f"{BASE_URL}/jobs",
        files={"file": ("test_documents.zip", zip_data, "application/zip")}
    )
    
    assert response.status_code == 202, f"Upload failed: {response.text}"
    job_data = response.json()
    job_id = job_data["job_id"]
    
    print(f"   ✓ Job created: {job_id}")
    print(f"   ✓ File count: {job_data['file_count']}")
    
    # Step 2: Poll for completion
    print("\n2. Waiting for conversion to complete...")
    max_attempts = 60  # 5 minutes maximum
    attempt = 0
    
    while attempt < max_attempts:
        response = requests.get(f"{BASE_URL}/jobs/{job_id}")
        assert response.status_code == 200
        
        status_data = response.json()
        status = status_data["status"]
        
        print(f"   Status: {status}", end="\r")
        
        if status == "COMPLETED":
            print("\n   ✓ Conversion completed!")
            break
        elif status == "FAILED":
            print("\n   ✗ Job failed!")
            print(f"   Details: {status_data}")
            return False
        
        time.sleep(5)
        attempt += 1
    
    if attempt >= max_attempts:
        print("\n   ✗ Timeout waiting for job completion")
        return False
    
    # Step 3: Check individual file statuses
    print("\n3. Checking individual file statuses:")
    for file_info in status_data["files"]:
        status_icon = "✓" if file_info["status"] == "COMPLETED" else "✗"
        print(f"   {status_icon} {file_info['filename']}: {file_info['status']}")
        if file_info["error_message"]:
            print(f"      Error: {file_info['error_message']}")
    
    # Step 4: Download results
    print("\n4. Downloading results...")
    response = requests.get(f"{BASE_URL}/jobs/{job_id}/download")
    
    assert response.status_code == 200, f"Download failed: {response.text}"
    assert response.headers["content-type"] == "application/zip"
    
    # Save the zip file
    output_file = f"test_output_{job_id}.zip"
    with open(output_file, "wb") as f:
        f.write(response.content)
    
    print(f"   ✓ Downloaded: {output_file}")
    print(f"   ✓ Size: {len(response.content)} bytes")
    
    # Step 5: Verify zip contents
    print("\n5. Verifying zip contents:")
    with zipfile.ZipFile(output_file, 'r') as zip_ref:
        file_list = zip_ref.namelist()
        for filename in file_list:
            file_info = zip_ref.getinfo(filename)
            print(f"   ✓ {filename} ({file_info.file_size} bytes)")
        
        assert len(file_list) == 3, "Expected 3 PDF files"
        assert all(f.endswith('.pdf') for f in file_list), "All files should be PDFs"
    
    print("\n" + "=" * 60)
    print("All tests passed! ✓")
    print("=" * 60)
    
    return True

if __name__ == "__main__":
    try:
        test_full_workflow()
    except Exception as e:
        print(f"\n✗ Test failed with error: {e}")
        import traceback
        traceback.print_exc()